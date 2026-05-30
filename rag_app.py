"""
MD Transaction RAG — FAISS 本地知识库主链路。

主链路：文档加载 → 文本清洗 → 数据脱敏 → 文本切片 → FAISS 检索 → Ollama 生成。
保留证据引用、召回日志和索引 manifest，用于定位来源与排查回答质量。
"""
from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Tuple

import requests
from langchain_core.embeddings import Embeddings

from masker import MASK_CONFIG, mask_documents

BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
FAISS_DIR = BASE_DIR / "faiss_index"
LOG_DIR = BASE_DIR / "logs"
TRACE_FILE = LOG_DIR / "rag_trace.jsonl"
MANIFEST_NAME = "manifest.json"

LLM_MODEL = os.getenv("RAG_LLM_MODEL", "qwen2.5:7b")
EMBED_MODEL = os.getenv("RAG_EMBED_MODEL", "nomic-embed-text")
OLLAMA_BASE_URL = os.getenv("RAG_OLLAMA_BASE_URL", "http://localhost:11434")
VECTOR_BACKEND = "faiss"

CHUNK_SIZE = int(os.getenv("RAG_CHUNK_SIZE", "650"))
CHUNK_OVERLAP = int(os.getenv("RAG_CHUNK_OVERLAP", "100"))
RETRIEVAL_K = int(os.getenv("RAG_RETRIEVAL_K", "10"))

PROMPT_VERSION = "rag_prompt_v4_qwen"
INGEST_VERSION = "2026-05-30-v4"


# ======================== Retrieval Balance Config ========================
# 先扩大候选召回，再做文档均衡，避免 700 页大 PDF 占满全部 top-k
CANDIDATE_K = int(os.getenv("RAG_CANDIDATE_K", "120"))
MAX_CHUNKS_PER_SOURCE = int(os.getenv("RAG_MAX_CHUNKS_PER_SOURCE", "5"))
# 文档权重：短文档/规则汇编更适合作为问答依据，给予轻微加权
# FAISS 分数通常是距离，越小越相似；后续用 score / weight 调整排序
DOC_WEIGHTS = {
    "蒙东新能源电站交易规则汇编_AI算法开发版.docx": 1.40,
    "蒙东电力交易市场交易员培训教材.pdf": 1.25,
    "20251021010756789.pdf": 1.00,
}

RAG_PROMPT = """\
你是”蒙东新能源电站交易规则”本地知识库问答助手。

回答规则：
1. 只能依据【参考片段】回答，不得编造未出现在片段中的规则、数值或流程。
   片段中的 Markdown 表格需逐列对应读取，不得混淆行列关系。
2. 涉及规则、时间、价格、电量、结算、主体条件时，必须在句末标注来源编号，如 [S1]、[S2]。
   数值必须与片段中的数字完全一致，不得四舍五入或近似。
3. 如果参考片段不足以回答，请明确说：”当前知识库未检索到足够依据”，并说明还需要查哪类文件。
4. 涉及”不得、不结转、除外、应当、必须、不得超过、按月、按日”等规则词时，不得反向改写原文含义。
5. 当多个来源的规则存在差异时，优先采信正式规则文件，并在注意事项中说明差异。
6. 条件性规则（”如...则...”）需完整保留条件和结论，不可省略前提。
7. 表达简洁专业，适合电力交易员阅读。

【参考片段】
{context}

【用户问题】
{question}

请按以下格式回答：
一、结论（简明直接回答用户问题）
二、依据（逐条列出支撑结论的规则原文，标注 [Sx]）
三、注意事项（规则适用的前置条件、例外情形、常见误解）
"""


# ======================== 0. Ollama Embeddings ========================
class OllamaEmbeddings(Embeddings):
    """
    用 Ollama HTTP API 做文本向量化，完全绕开 PyTorch / sentence-transformers。
    需要本地已拉取：ollama pull nomic-embed-text
    """

    def __init__(self, model: str = EMBED_MODEL, base_url: str = OLLAMA_BASE_URL):
        self.model = model
        self.base_url = base_url.rstrip("/")

    def _embed(self, texts: List[str]) -> List[List[float]]:
        vectors = []
        for text in texts:
            resp = requests.post(
                f"{self.base_url}/api/embeddings",
                json={"model": self.model, "prompt": text},
                timeout=60,
            )
            resp.raise_for_status()
            vectors.append(resp.json()["embedding"])
        return vectors

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        print(f"  encoding {len(texts)} texts via Ollama ({self.model})...")
        return self._embed(texts)

    def embed_query(self, text: str) -> List[float]:
        return self._embed([text])[0]

    def __call__(self, text: str) -> List[float]:
        return self.embed_query(text)


# ======================== 1. Load ========================
def _table_to_markdown(table):
    """将二维表格 list-of-lists 转为 Markdown 表格字符串。

    pdfplumber / python-docx 提取的表格行内可能有 None、换行符和空格，
    这里统一清洗为紧凑的 Markdown 表格块，便于 LLM 准确理解行列关系。
    """
    if not table:
        return ""
    cleaned = []
    for row in table:
        cells = [(cell or "").strip().replace("\n", " ").replace("|", "\\|") for cell in row]
        cleaned.append(cells)
    cleaned = [row for row in cleaned if any(c for c in row)]
    if not cleaned:
        return ""
    col_count = max(len(row) for row in cleaned)
    for row in cleaned:
        while len(row) < col_count:
            row.append("")
    lines = []
    lines.append("| " + " | ".join(cleaned[0]) + " |")
    lines.append("|" + "|".join([" --- " for _ in range(col_count)]) + "|")
    for row in cleaned[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _load_pdf_pages(fp):
    """使用 pdfplumber 逐页解析 PDF，自动检测表格并输出 Markdown。

    返回 List[Document]，每页一个 Document。
    若 pdfplumber 打不开文件，自动降级为 PyPDFLoader。
    """
    from langchain_core.documents import Document

    docs = []
    try:
        import pdfplumber
        pdf = pdfplumber.open(str(fp))
    except Exception:
        print(f"    [WARN] pdfplumber 无法打开 {fp.name}，降级为 PyPDFLoader")
        from langchain_community.document_loaders import PyPDFLoader
        return PyPDFLoader(str(fp)).load()

    for page_num, page in enumerate(pdf.pages):
        text = page.extract_text() or ""

        tables = page.extract_tables()
        md_blocks = [_table_to_markdown(t) for t in tables if _table_to_markdown(t)]
        if md_blocks:
            text += "\n\n" + "\n\n".join(md_blocks)

        if not text.strip():
            continue

        doc = Document(page_content=text)
        doc.metadata["page"] = page_num
        doc.metadata["has_tables"] = len(md_blocks) > 0
        doc.metadata["table_count"] = len(md_blocks)
        docs.append(doc)

    pdf.close()
    return docs


def _load_docx_sections(fp):
    """使用 python-docx 解析 DOCX，保留段落-表格交错顺序，表格转 Markdown。

    返回 List[Document]（当前仅含一个 Document，后续由 splitter 切片）。
    若 python-docx 不可用，自动降级为 Docx2txtLoader。
    """
    from langchain_core.documents import Document

    try:
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn
    except ImportError:
        print(f"    [WARN] python-docx 不可用，降级为 Docx2txtLoader")
        from langchain_community.document_loaders import Docx2txtLoader
        return Docx2txtLoader(str(fp)).load()

    doc = DocxDocument(str(fp))

    para_by_elem = {p._element: p for p in doc.paragraphs}
    table_by_elem = {t._element: t for t in doc.tables}

    parts = []

    for child in doc.element.body:
        tag = child.tag
        if tag == qn("w:p"):
            para = para_by_elem.get(child)
            if para:
                text = para.text
                if text and text.strip():
                    parts.append(("text", text.strip()))
        elif tag == qn("w:tbl"):
            table = table_by_elem.get(child)
            if table:
                rows = []
                for row in table.rows:
                    cells = [
                        cell.text.strip().replace("\n", " ").replace("|", "\\|")
                        for cell in row.cells
                    ]
                    rows.append(cells)
                if rows:
                    parts.append(("table", rows))

    blocks = []
    table_count = 0
    for ptype, content in parts:
        if ptype == "text":
            blocks.append(content)
        elif ptype == "table":
            md = _table_to_markdown(content)
            if md:
                blocks.append(md)
                table_count += 1

    if not blocks:
        return []

    full_text = "\n\n".join(blocks)
    out = Document(page_content=full_text)
    out.metadata["has_tables"] = table_count > 0
    out.metadata["table_count"] = table_count
    return [out]


def load_documents(data_dir: Path):
    """从 data_dir 加载全部 PDF 和 DOCX，为每个文件附加 filename/doc_type/source metadata。"""
    docs = []
    if not data_dir.exists():
        print(f"ERROR: data directory not found: {data_dir}")
        return docs

    for fp in sorted(data_dir.glob("*.pdf")):
        print(f"  [PDF] {fp.name}")
        loaded = _load_pdf_pages(fp)
        for d in loaded:
            d.metadata["filename"] = fp.name
            d.metadata["doc_type"] = "pdf"
            d.metadata["source"] = str(fp.relative_to(BASE_DIR))
        docs.extend(loaded)
        table_pages = sum(1 for d in loaded if d.metadata.get("has_tables"))
        if table_pages:
            print(f"         {len(loaded)} pages, {table_pages} with table(s)")

    for fp in sorted(data_dir.glob("*.docx")):
        print(f"  [DOCX] {fp.name}")
        loaded = _load_docx_sections(fp)
        for d in loaded:
            d.metadata["filename"] = fp.name
            d.metadata["doc_type"] = "docx"
            d.metadata["source"] = str(fp.relative_to(BASE_DIR))
        docs.extend(loaded)
        if loaded and loaded[0].metadata.get("has_tables"):
            print(f"         {loaded[0].metadata.get('table_count', 0)} table(s) → Markdown")

    print(f"  loaded {len(docs)} pages/sections")
    return docs


# ======================== 1.5 Clean + Index Manifest ========================
CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def clean_text(text: str) -> str:
    """清理 PDF 抽取噪声，避免控制字符污染 embedding 和 prompt。"""
    if not text:
        return ""
    text = CONTROL_CHARS_RE.sub(" ", text)
    text = text.replace("\ufffd", " ")
    lines = []
    for raw in text.splitlines():
        line = re.sub(r"[ \t]+", " ", raw).strip()
        if line:
            lines.append(line)
    return "\n".join(lines)


def clean_documents(docs):
    changed = 0
    for doc in docs:
        original = doc.page_content
        cleaned = clean_text(original)
        if cleaned != original:
            changed += 1
            doc.page_content = cleaned
            doc.metadata["cleaned"] = True
    if changed:
        print(f"  cleaned text noise in {changed} pages/sections")
    return docs


# ======================== 1.6 入库前过滤低价值页面 ========================
_LOW_VALUE_SIGNATURES = [
    "目录", "目  录", "CONTENTS", "Table of Contents",
    "前言", "序言", "绪论", "引言",
    "致谢", "鸣谢", "ACKNOWLEDGEMENTS",
    "参考文献", "参考资料", "References",
    "附录", "Appendix",
    "声明", "免责声明", "Disclaimer",
]


def _is_low_value_page(text: str) -> bool:
    """判断整个页面是否为低价值页（封面、目录、致谢等），不入库。"""
    clean = text.strip()
    if len(clean) < 30:
        return True
    first_line = clean.split("\n")[0].strip()
    for sig in _LOW_VALUE_SIGNATURES:
        if sig in first_line:
            return True
    # 纯考题页特征：大量 A. B. C. D. 选项
    option_matches = len(re.findall(r"[A-D][．.。、]\s*\S", clean))
    if option_matches >= 6:
        return True
    # 纯数字/页码行
    num_blank_ratio = sum(1 for c in clean if c.isdigit() or c in ".。 　\t\r") / max(len(clean), 1)
    if num_blank_ratio > 0.85 and len(clean) < 200:
        return True
    return False


def filter_documents(docs):
    """过滤封面、目录、致谢、考题等低价值页面。"""
    kept, removed = [], 0
    for doc in docs:
        if _is_low_value_page(doc.page_content):
            removed += 1
        else:
            kept.append(doc)
    if removed:
        print(f"  filtered {removed} low-value pages (cover/TOC/preface/exam)")
    return kept


def _relative_path(path: Path) -> str:
    try:
        return str(path.relative_to(BASE_DIR))
    except ValueError:
        return str(path)


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _data_files_manifest() -> List[Dict[str, Any]]:
    files = sorted(DATA_DIR.glob("*.pdf")) + sorted(DATA_DIR.glob("*.docx"))
    return [
        {
            "path": _relative_path(fp),
            "size": fp.stat().st_size,
            "mtime_ns": fp.stat().st_mtime_ns,
            "sha256": _sha256_file(fp),
        }
        for fp in files
    ]


def _mask_config_hash() -> str:
    payload = json.dumps(MASK_CONFIG, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:12]


def current_index_manifest() -> Dict[str, Any]:
    return {
        "ingest_version": INGEST_VERSION,
        "vector_backend": VECTOR_BACKEND,
        "embed_model": EMBED_MODEL,
        "chunk_size": CHUNK_SIZE,
        "chunk_overlap": CHUNK_OVERLAP,
        "mask_config_hash": _mask_config_hash(),
        "data_files": _data_files_manifest(),
    }


def _manifest_path(index_dir: Path) -> Path:
    return index_dir / MANIFEST_NAME


def write_index_manifest(index_dir: Path):
    index_dir.mkdir(exist_ok=True)
    with _manifest_path(index_dir).open("w", encoding="utf-8") as f:
        json.dump(current_index_manifest(), f, ensure_ascii=False, indent=2)


def _manifest_diff(saved: Dict[str, Any], current: Dict[str, Any]) -> List[str]:
    changed = []
    for key in ("ingest_version", "vector_backend", "embed_model", "chunk_size", "chunk_overlap", "mask_config_hash"):
        if saved.get(key) != current.get(key):
            changed.append(key)
    saved_files = {item.get("path"): item for item in saved.get("data_files", [])}
    current_files = {item.get("path"): item for item in current.get("data_files", [])}
    if saved_files != current_files:
        changed.append("data_files")
    return changed


def warn_if_index_stale(index_dir: Path):
    manifest_file = _manifest_path(index_dir)
    if not manifest_file.exists():
        print("  [WARN] cached index has no manifest; rebuild once after optimization to enable freshness checks.")
        return
    try:
        saved = json.loads(manifest_file.read_text(encoding="utf-8"))
        changed = _manifest_diff(saved, current_index_manifest())
    except Exception as e:
        print(f"  [WARN] failed to read index manifest: {e}")
        return
    if changed:
        print(f"  [WARN] cached index may be stale; changed: {', '.join(changed)}")


# ======================== 2. Split + Metadata ========================
def _safe_slug(text: str) -> str:
    stem = Path(text).stem if text else "doc"
    keep = []
    for ch in stem:
        if ch.isalnum() or "\u4e00" <= ch <= "\u9fff":
            keep.append(ch)
        else:
            keep.append("_")
    return "".join(keep).strip("_")[:40] or "doc"


def _hash_text(text: str) -> str:
    return hashlib.md5(text.encode("utf-8", errors="ignore")).hexdigest()[:8]


def _guess_section_title(text: str) -> str:
    """提取章/节/条标题，支持多级标题串联以定位原文位置。

    优先级：第X章 → 第X节 → 第X条 → 中文编号 → 业务关键词。
    多级命中时以 "  ▸  " 连接，例如 "第三章 市场注册  ▸  第二节 注册流程"。
    """
    found = []
    seen_chapter = False

    for raw in text.splitlines()[:10]:
        line = raw.strip()
        if not line:
            continue
        if not seen_chapter:
            m = re.match(r"第[一二三四五六七八九十百千]+章[^\n]{0,60}", line)
            if m:
                found.append(m.group())
                seen_chapter = True
                continue
        m = re.match(r"第[一二三四五六七八九十百千]+节[^\n]{0,60}", line)
        if m:
            found.append(m.group())
            continue
        m = re.match(r"第[一二三四五六七八九十百千]+条[^\n]{0,80}", line)
        if m:
            found.append(m.group())
            continue

    if found:
        return "  ▸  ".join(found)

    # 回退：中文数字编号（一、二、...）或关键词行
    for raw in text.splitlines()[:5]:
        line = raw.strip()
        if not line:
            continue
        if re.match(r"[一二三四五六七八九十]+[、．.]", line) and len(line) <= 50:
            return line
        if len(line) <= 35 and any(
            key in line for key in (
                "市场", "交易", "结算", "注册", "申报", "价格", "保量",
                "中长期", "现货", "日前", "日内", "合约", "电量", "电价",
                "规则", "流程", "条件", "方式", "机制", "主体",
            )
        ):
            return line
    return ""


def enrich_chunk_metadata(chunks):
    counters: Dict[str, int] = {}
    for chunk in chunks:
        meta = chunk.metadata or {}
        source = meta.get("source", "unknown")
        filename = meta.get("filename") or Path(source).name
        doc_id = _safe_slug(filename)
        page = meta.get("page", "")
        page_label = str(meta.get("page_label") or (page + 1 if isinstance(page, int) else page or ""))

        key = f"{doc_id}_p{page_label}"
        counters[key] = counters.get(key, 0) + 1
        chunk_index = counters[key]
        text_hash = _hash_text(chunk.page_content)
        chunk_id = f"{doc_id}_p{page_label}_{chunk_index:03d}_{text_hash}"

        meta.update(
            {
                "chunk_id": chunk_id,
                "doc_id": doc_id,
                "filename": filename,
                "source": source,
                "page_label": page_label,
                "section_title": _guess_section_title(chunk.page_content),
                "chunk_index": chunk_index,
                "text_hash": text_hash,
                "masked": True,
                "mask_policy": "business_numbers_preserved_v2",
                "ingest_version": INGEST_VERSION,
            }
        )
        chunk.metadata = meta
    return chunks


# ======================== 表格保护 & chunk 前缀 ========================
_TABLE_RE = re.compile(
    r"(\|[^\n]+\|\n\|[-\s|:]+\|\n(?:\|[^\n]*\|\n?)+)",
    re.MULTILINE,
)


def _protect_tables(text: str):
    """提取 Markdown 表格，替换为占位符，避免切分器拆散表格。"""
    tables = []

    def _replace(m):
        tables.append(m.group(1))
        return f"\n\n<!--TBL_{len(tables) - 1}-->\n\n"

    protected = _TABLE_RE.sub(_replace, text)
    return protected, tables


def _restore_tables(chunks, all_tables):
    """将占位符还原为 Markdown 表格。"""
    for chunk in chunks:
        text = chunk.page_content
        for i, tbl in enumerate(all_tables):
            text = text.replace(f"<!--TBL_{i}-->", tbl)
        chunk.page_content = text


def _build_chunk_prefix(meta: dict) -> str:
    """构建 chunk 上下文前缀，写入 embedding 以区分不同文档的相似段落。"""
    filename = meta.get("filename", "")
    section = meta.get("section_title", "")
    parts = []
    if filename:
        parts.append(f"[{filename}]")
    if section:
        parts.append(f"[{section}]")
    return " ".join(parts) + "\n" if parts else ""


def split_documents(docs):
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=[
            # 文档结构边界（优先在此断开）
            "\n第",           # 第X章/节/条
            "\n# ", "\n## ",  # Markdown 标题
            "\n一、", "\n二、", "\n三、", "\n四、", "\n五、",
            "\n六、", "\n七、", "\n八、", "\n九、", "\n十、",
            # 自然段落
            "\n\n", "\n",
            # 句子边界
            "。", "；", "！", "？",
            # 英文
            ". ", "; ",
            # 最后手段
            " ", "",
        ],
    )

    # 1. 保护表格：提取后暂存，切完再还原
    all_tables = []
    for doc in docs:
        doc.page_content, tables = _protect_tables(doc.page_content)
        all_tables.extend(tables)

    # 2. 切分
    chunks = splitter.split_documents(docs)

    # 3. 还原表格
    _restore_tables(chunks, all_tables)

    # 4. 元数据增强
    chunks = enrich_chunk_metadata(chunks)

    # 5. 添加文档/章节上下文前缀（仅用于 embedding，不修改 page_content 的后续使用）
    #    前缀帮助 FAISS 区分不同文档中语义相似但来源不同的片段
    for chunk in chunks:
        prefix = _build_chunk_prefix(chunk.metadata)
        chunk.page_content = prefix + chunk.page_content

    print(f"  {len(chunks)} chunks")
    return chunks


# ======================== 3. Vectorstore ========================
def build_faiss_vectorstore(chunks):
    from langchain_community.vectorstores import FAISS

    embeddings = OllamaEmbeddings()
    print(f"  embedding via Ollama: {EMBED_MODEL}")
    print(f"  building FAISS -> {FAISS_DIR}")
    vs = FAISS.from_documents(chunks, embeddings)
    if FAISS_DIR.exists():
        shutil.rmtree(FAISS_DIR)
    vs.save_local(str(FAISS_DIR))
    write_index_manifest(FAISS_DIR)
    return vs


def load_faiss_vectorstore():
    from langchain_community.vectorstores import FAISS

    index_file = FAISS_DIR / "index.faiss"
    pkl_file = FAISS_DIR / "index.pkl"
    if not index_file.exists() or not pkl_file.exists():
        return None
    warn_if_index_stale(FAISS_DIR)
    embeddings = OllamaEmbeddings()
    return FAISS.load_local(
        str(FAISS_DIR),
        embeddings,
        allow_dangerous_deserialization=True,
    )


def build_vectorstore(chunks):
    return build_faiss_vectorstore(chunks)


def load_vectorstore():
    return load_faiss_vectorstore()


# ======================== 4. Ask + Sources + Trace ========================
def _short_excerpt(text: str, limit: int = 150) -> str:
    text = " ".join((text or "").split())
    return text[:limit] + ("..." if len(text) > limit else "")


def _format_context(results: List[Tuple[Any, float]]) -> Tuple[str, List[Dict[str, Any]]]:
    context_blocks = []
    sources = []

    for rank, (doc, score) in enumerate(results, start=1):
        meta = doc.metadata or {}
        sid = f"S{rank}"
        filename = meta.get("filename") or Path(meta.get("source", "unknown")).name
        page_label = meta.get("page_label") or meta.get("page") or ""
        section_title = meta.get("section_title") or ""
        chunk_id = meta.get("chunk_id") or f"chunk_{rank}"
        excerpt = _short_excerpt(doc.page_content)

        # 构建人可读的定位串，前端据此在原文件中定位
        loc_parts = []
        if section_title:
            loc_parts.append(section_title)
        if page_label not in (None, ""):
            loc_parts.append(f"p.{page_label}")
        location_text = "  |  ".join(loc_parts) if loc_parts else filename

        sources.append(
            {
                "id": sid,
                "rank": rank,
                "score": float(score) if score is not None else None,
                "filename": filename,
                "source": meta.get("source", "unknown"),
                "page_label": page_label,
                "section_title": section_title,
                "location_text": location_text,
                "chunk_id": chunk_id,
                "excerpt": excerpt,
            }
        )

        context_blocks.append(
            f"[{sid}] 文件：{filename} | 定位：{location_text} | chunk_id：{chunk_id}\n"
            f"{doc.page_content}"
        )

    return "\n\n---\n\n".join(context_blocks), sources


def _call_ollama(prompt_text: str) -> str:
    resp = requests.post(
        f"{OLLAMA_BASE_URL.rstrip('/')}/api/generate",
        json={
            "model": LLM_MODEL,
            "prompt": prompt_text,
            "stream": False,
            "options": {
                "temperature": 0.1,
                "top_p": 0.9,
                "num_ctx": 8192,
            },
        },
        timeout=180,
    )
    resp.raise_for_status()
    return resp.json().get("response", "").strip()


def _write_trace(record: Dict[str, Any]):
    LOG_DIR.mkdir(exist_ok=True)
    with TRACE_FILE.open("a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False) + "\n")


def _search_with_scores(vectorstore, question: str, k: int) -> List[Tuple[Any, Any]]:
    if hasattr(vectorstore, "similarity_search_with_score"):
        return vectorstore.similarity_search_with_score(question, k=k)
    docs = vectorstore.similarity_search(question, k=k)
    return [(doc, None) for doc in docs]


# ======================== Reranker helpers ========================
def _keyword_relevance(query: str, text: str) -> float:
    """计算 query 与 chunk 的关键词重叠度，用于弥补 FAISS 纯语义检索的不足。

    提取 query 中的 2-4 字中文词和英文词，统计在 chunk 中的命中率。
    返回 0-1 之间的分数。
    """
    text_lower = text.lower()
    # 提取中文二字及以上词组 + 英文单词
    terms = re.findall(r"[一-鿿]{2,4}|[a-zA-Z]{2,}", query)
    if not terms:
        return 0.5
    hits = sum(1 for t in terms if t.lower() in text_lower)
    return hits / len(terms)


def _is_noise_chunk(text: str) -> bool:
    """检测低价值 chunk（目录、封面、考题、过短纯数字等），返回 True 表示应降权。"""
    clean = text.strip()
    if len(clean) < 60:
        return True
    first_line = clean.split("\n")[0].strip()
    # 目录 / 前言 / 致谢 等
    if any(kw in first_line for kw in ("目录", "目  录", "CONTENTS", "前言", "致谢", "参考文献")):
        return True
    # 大量点线连接（目录特征）
    dots = clean.count(".") + clean.count("…") + clean.count("．．")
    if dots > 20 and len(clean) < 500:
        return True
    # 纯数字/空格占比过高（疑似损坏的表格或索引）
    num_blank = sum(1 for c in clean if c.isdigit() or c in ".。 　\t")
    if num_blank / max(len(clean), 1) > 0.7:
        return True
    return False


def _retrieve(vectorstore, question: str) -> List[Tuple[Any, Any]]:
    """文档均衡检索 + 关键词/噪音重排增强。

    保留原始固定参数检索骨架，仅在上层叠加两层增强：
    - 关键词命中率微调排序（最多影响 30%）
    - 目录/封面/碎片等低价值 chunk 自动降权

    其余不变：CANDIDATE_K 初筛 → DOC_WEIGHTS → 来源均衡 → RETRIEVAL_K 定长输出。
    """
    candidate_k = max(CANDIDATE_K, RETRIEVAL_K)

    raw_results = _search_with_scores(vectorstore, question, candidate_k)

    def _source_name(doc):
        meta = doc.metadata or {}
        src = meta.get("source") or meta.get("filename") or "unknown"
        src = str(src)
        try:
            return Path(src).name
        except Exception:
            return src

    def _weight_for_source(source_name):
        for key, weight in DOC_WEIGHTS.items():
            if key in source_name:
                return float(weight)
        return 1.0

    ranked = []
    for rank, (doc, score) in enumerate(raw_results, start=1):
        source_name = _source_name(doc)
        weight = _weight_for_source(source_name)

        try:
            numeric_score = float(score) if score is not None else float(rank)
        except Exception:
            numeric_score = float(rank)

        # FAISS 距离，越小越相似。高权重文档除以 weight 使其排序更靠前。
        adjusted_score = numeric_score / max(weight, 0.01)

        # ---- 增强层 1：关键词重排（微调，最多影响 30%） ----
        kw_rel = _keyword_relevance(question, doc.page_content)
        adjusted_score *= (1.0 - kw_rel * 0.30)

        # ---- 增强层 2：低价值 chunk 降权 ----
        if _is_noise_chunk(doc.page_content):
            adjusted_score *= 2.5

        ranked.append({
            "doc": doc,
            "score": score,
            "source": source_name,
            "adjusted_score": adjusted_score,
        })

    ranked.sort(key=lambda x: x["adjusted_score"])

    selected = []
    source_counts = {}

    # 第一轮：每个文档最多保留 MAX_CHUNKS_PER_SOURCE 个
    for item in ranked:
        src = item["source"]
        if source_counts.get(src, 0) >= MAX_CHUNKS_PER_SOURCE:
            continue

        selected.append((item["doc"], item["score"]))
        source_counts[src] = source_counts.get(src, 0) + 1

        if len(selected) >= RETRIEVAL_K:
            break

    # 第二轮：如果不够 top-k，则放宽限制补齐
    if len(selected) < RETRIEVAL_K:
        selected_ids = {id(doc) for doc, _ in selected}

        for item in ranked:
            doc = item["doc"]
            if id(doc) in selected_ids:
                continue

            selected.append((doc, item["score"]))
            selected_ids.add(id(doc))

            if len(selected) >= RETRIEVAL_K:
                break

    # 打印召回来源分布
    try:
        dist = {}
        for doc, _ in selected:
            src = _source_name(doc)
            dist[src] = dist.get(src, 0) + 1

        print("  召回来源分布: " + " | ".join(f"{k}: {v}" for k, v in dist.items()))
    except Exception:
        pass

    return selected


# ======================== Query 改写 ========================
_QUERY_EXPANSIONS = {
    "结算": "结算方式 结算周期 结算电量 结算价格 电费结算",
    "注册": "市场注册 注册流程 注册条件 注册材料 市场主体注册",
    "申报": "申报流程 申报时间 申报要求 申报材料 交易申报",
    "中长期": "中长期合约 中长期交易 中长期结算 年度合约 月度合约",
    "现货": "现货交易 现货市场 日前市场 日内市场 实时市场",
    "价格": "电价 限价 出清价格 合约价格 度电价格",
    "电量": "交易电量 结算电量 申报电量 保量保价",
    "考核": "偏差考核 免考核 考核标准 考核费用",
    "合同": "入市协议 购售电合同 输配电合同",
    "偏差": "偏差电量 偏差考核 偏差结算 正偏差 负偏差",
}


def _rewrite_query(question: str) -> str:
    """将用户短问题扩展为信息更丰富的检索 query。

    规则：
    - 问题已 >= 25 字符 → 直接使用原问题（信息量足够）
    - 短问题 → 匹配关键词词典做术语扩展
    - 不调用 LLM，零延迟
    """
    if len(question) >= 25:
        return question
    terms_added = set()
    for key, expansion in _QUERY_EXPANSIONS.items():
        if key in question:
            for term in expansion.split():
                if term not in question and term not in terms_added:
                    terms_added.add(term)
    if terms_added:
        return question + " " + " ".join(terms_added)
    return question


# ======================== 答案校验 ========================
_NEGATION_PAIRS = [
    ("不得", "可以"),
    ("不应", "应该"),
    ("不结转", "结转"),
    ("禁止", "允许"),
    ("除外", "包括"),
    ("不予", "予以"),
    ("不得参与", "可以参与"),
    ("不能", "可以"),
]


def _verify_answer(answer: str, sources: list) -> str:
    """对 LLM 输出做方向性校验，防止否定词反转。

    检查来源文本中出现的否定/限制词是否在回答中被反向改写。
    如果发现可疑反转，在回答末尾追加校验提示。
    """
    if answer.startswith("[ERROR]"):
        return answer

    source_text = " ".join(s.get("excerpt", "") for s in sources)
    warnings = []

    for prohibited, allowed in _NEGATION_PAIRS:
        if prohibited in source_text:
            if allowed in answer and prohibited not in answer:
                warnings.append(f"原文含「{prohibited}」，回答中未出现，请以原文为准")

    if warnings:
        answer += "\n\n---\n⚠️ **自动校验提示：**\n"
        for w in warnings[:3]:
            answer += f"- {w}\n"

    return answer


def ask(vectorstore, question: str) -> Dict[str, Any]:
    # 检索用改写后的 query，LLM 用原始问题
    search_query = _rewrite_query(question)
    if search_query != question:
        print(f"  query rewritten: {question[:40]}... → {search_query[:60]}...")

    results = _retrieve(vectorstore, search_query)
    context, sources = _format_context(results)
    prompt_text = RAG_PROMPT.format(context=context, question=question)

    try:
        answer = _call_ollama(prompt_text)
    except requests.ConnectionError:
        answer = "[ERROR] 无法连接 Ollama，请先运行: ollama serve"
    except requests.Timeout:
        answer = "[ERROR] 模型推理超时，请检查 Ollama 是否正常运行，或换用更小 top-k。"
    except requests.HTTPError as e:
        answer = f"[ERROR] Ollama HTTP 调用失败：{e}"

    answer = _verify_answer(answer, sources)

    result = {
        "ts": datetime.now().isoformat(timespec="seconds"),
        "question": question,
        "answer": answer,
        "sources": sources,
        "model": LLM_MODEL,
        "embed_model": EMBED_MODEL,
        "vector_backend": VECTOR_BACKEND,
        "retrieval_k": RETRIEVAL_K,
        "candidate_k": CANDIDATE_K,
        "max_chunks_per_source": MAX_CHUNKS_PER_SOURCE,
        "prompt_version": PROMPT_VERSION,
    }
    _write_trace(result)
    return result


def print_answer(result: Dict[str, Any]):
    print(result["answer"])
    print()
    print("参考片段：")
    for s in result.get("sources", []):
        score_text = f"score={s['score']:.4f}" if isinstance(s.get("score"), float) else "score=N/A"
        loc = f"p.{s['page_label']}" if s.get("page_label") not in (None, "") else "p.N/A"
        section = f" | {s['section_title']}" if s.get("section_title") else ""
        print(f"[{s['id']}] {s['filename']} | {loc} | {score_text}{section}")
        print(f"    chunk_id={s['chunk_id']}")
        print(f"    摘录：{s['excerpt']}")
    print(f"\n[trace] {TRACE_FILE}")


# ======================== 5. CLI ========================
def interactive_qa(vectorstore):
    print()
    print("=" * 72)
    print(f"  MD Transaction RAG | backend: {VECTOR_BACKEND} | model: {LLM_MODEL} | top-k: {RETRIEVAL_K}")
    print("  输入问题直接回车；输入 quit / exit / q 退出")
    print("=" * 72)
    print()

    while True:
        try:
            q = input(">>> ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\n  bye.")
            break

        if not q:
            continue
        if q.lower() in ("quit", "exit", "q"):
            print("  bye.")
            break

        print("  检索并生成中...\n")
        try:
            result = ask(vectorstore, q)
            print_answer(result)
            print()
        except Exception:
            import traceback
            print("-" * 40)
            traceback.print_exc()
            print("-" * 40)
            print()


# ======================== main ========================
def main():
    print("=" * 72)
    print("  MD Transaction - RAG Knowledge Base")
    print("=" * 72)

    vs = load_vectorstore()
    if vs is not None:
        print(f"\n  Using cached vectorstore: {FAISS_DIR}")
        print("  如修改了知识源或脱敏配置，请删除索引目录后重建。\n")
    else:
        print("\n[1/3] Loading documents...")
        docs = load_documents(DATA_DIR)
        if not docs:
            print("ERROR: no documents found in data/")
            return

        print("\n[*] Cleaning extracted text...")
        docs = clean_documents(docs)

        print("\n[*] Filtering low-value pages...")
        docs = filter_documents(docs)

        print("\n[*] Masking sensitive data...")
        docs = mask_documents(docs, config=MASK_CONFIG)

        print("\n[2/3] Splitting and enriching metadata...")
        chunks = split_documents(docs)

        print("\n[3/3] Building vectorstore...")
        vs = build_vectorstore(chunks)
        print("  done.\n")

    print("[OK] Ready.\n")
    interactive_qa(vs)


if __name__ == "__main__":
    main()
